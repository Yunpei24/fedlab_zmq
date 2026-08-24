from __future__ import annotations

from pathlib import Path
from math import sin, cos, pi

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/joshuajusteyunpeinikiema/Documents/PhD_UM6P/Stat_of_art_EnergyEfficientFL/fedlab_zmq")
OUT = ROOT / "docs" / "FedBCR_DP_cadre_de_recherche.docx"
ASSET_DIR = ROOT / "docs" / "_fedbcr_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

NAVY = "16324F"
BLUE = "2E74B5"
MID_BLUE = "4E8BC4"
LIGHT_BLUE = "EAF2F8"
PALE = "F4F6F9"
GREEN = "2F7D68"
LIGHT_GREEN = "E7F3EF"
ORANGE = "C87533"
LIGHT_ORANGE = "FAEEDF"
RED = "A94442"
LIGHT_RED = "F8E8E8"
GRAY = "5B6573"
LIGHT_GRAY = "E9EDF2"
WHITE = "FFFFFF"
BLACK = "1A1A1A"


def rgb(hexstr: str) -> RGBColor:
    return RGBColor.from_string(hexstr)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, twips: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths_twips):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_twips):
            set_cell_width(cell, width)
            set_cell_margins(cell)


def set_repeat_header(p):
    p_pr = p._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def add_field(run, instruction: str):
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_repeat_and_keep(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    for tag in ("keepNext", "keepLines"):
        if p_pr.find(qn(f"w:{tag}")) is None:
            p_pr.append(OxmlElement(f"w:{tag}"))


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def get_font(size=30, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            pass
    return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline=BLUE, radius=24, width=4):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_arrow(draw, p1, p2, fill=BLUE, width=7, head=18):
    draw.line((p1, p2), fill=fill, width=width)
    x1, y1 = p1
    x2, y2 = p2
    angle = __import__("math").atan2(y2 - y1, x2 - x1)
    a1 = angle + 2.6
    a2 = angle - 2.6
    draw.polygon([(x2, y2), (x2 + head * cos(a1), y2 + head * sin(a1)), (x2 + head * cos(a2), y2 + head * sin(a2))], fill=fill)


def make_geometry_figure(path: Path):
    w, h = 1800, 980
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title = get_font(42, True)
    body = get_font(30)
    small = get_font(25)
    d.text((65, 40), "D’un point-sonde à un descripteur local orienté de frontière", font=title, fill="#16324F")
    d.line((115, 820, 1640, 820), fill="#9AA5B1", width=4)
    d.line((190, 890, 190, 150), fill="#9AA5B1", width=4)
    pts = []
    for i in range(920):
        x = 300 + i * 1.35
        y = 500 + 120 * sin(i / 155) - 0.07 * i
        pts.append((x, y))
    d.line(pts, fill="#2E74B5", width=11)
    d.text((1020, 245), "frontière locale  Gamma_k,ab : m_k,ab(z)=0", font=body, fill="#2E74B5")
    zq = (730, 710)
    bq = (774, 532)
    d.ellipse((zq[0]-13, zq[1]-13, zq[0]+13, zq[1]+13), fill="#C87533")
    d.ellipse((bq[0]-14, bq[1]-14, bq[0]+14, bq[1]+14), fill="#2F7D68")
    d.text((505, 735), "sonde publique z_q", font=body, fill="#C87533")
    d.text((795, 500), "projection b_k,q", font=body, fill="#2F7D68")
    draw_arrow(d, zq, bq, fill="#C87533", width=8, head=22)
    d.text((755, 625), "delta_k,q  n_k,q", font=body, fill="#C87533")
    n_end = (865, 390)
    draw_arrow(d, bq, n_end, fill="#2F7D68", width=8, head=22)
    d.text((875, 355), "normale unitaire n_k,q", font=body, fill="#2F7D68")
    rounded_box(d, (1035, 520, 1690, 850), "#F4F6F9", outline="#5B6573", radius=20, width=3)
    d.text((1080, 555), "Descripteur transmis", font=get_font(31, True), fill="#16324F")
    items = ["• identifiant de la sonde q", "• paire de classes (a,b)", "• décalage signé delta_k,q", "• normale n_k,q", "• résidu, rayon de validité, poids"]
    yy = 620
    for item in items:
        d.text((1090, yy), item, font=small, fill="#1A1A1A")
        yy += 43
    d.text((65, 910), "Le point privé n’est pas envoyé : la sonde est publique/synthétique et la géométrie est ensuite agrégée sous confidentialité.", font=small, fill="#5B6573")
    img.save(path, dpi=(180, 180))


def make_pipeline_figure(path: Path):
    w, h = 2200, 1040
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 35), "FedBCR-DP : changer l’objet de collaboration", font=get_font(46, True), fill="#16324F")
    d.text((72, 100), "Les clients ne fusionnent pas leurs paramètres : ils décrivent localement une frontière, puis le serveur reconstruit un champ global.", font=get_font(28), fill="#5B6573")
    boxes = [
        ("1", "Repère commun", "encodeur E gelé\n+ sondes publiques"),
        ("2", "Géométrie locale", "marge, normale,\nprojection, confiance"),
        ("3", "Esquisse privée", "base additive,\nclipping + bruit partagé"),
        ("4", "SecAgg", "seule la somme\nbruitée est révélée"),
        ("5", "Reconstruction", "SDF neuronale +\nEikonal + UOT/topologie"),
        ("6", "Modèle utilisable", "classifieur global\nou segmentation"),
    ]
    colors = ["#EAF2F8", "#E7F3EF", "#FAEEDF", "#F8E8E8", "#EAF2F8", "#E7F3EF"]
    x0, bw, gap, y0, bh = 70, 315, 40, 250, 400
    for i, ((num, title, text), fill) in enumerate(zip(boxes, colors)):
        x = x0 + i * (bw + gap)
        rounded_box(d, (x, y0, x+bw, y0+bh), fill, outline="#2E74B5", radius=26, width=4)
        d.ellipse((x+18, y0+18, x+76, y0+76), fill="#16324F")
        d.text((x+37, y0+29), num, font=get_font(27, True), fill="white", anchor="mm")
        d.text((x+bw/2, y0+118), title, font=get_font(29, True), fill="#16324F", anchor="mm")
        lines = text.split("\n")
        for j, line in enumerate(lines):
            d.text((x+bw/2, y0+205+j*53), line, font=get_font(25), fill="#1A1A1A", anchor="mm")
        if i < len(boxes)-1:
            draw_arrow(d, (x+bw+7, y0+bh/2), (x+bw+gap-7, y0+bh/2), fill="#4E8BC4", width=6, head=17)
    rounded_box(d, (115, 750, 2085, 950), "#F4F6F9", outline="#5B6573", radius=20, width=3)
    d.text((165, 790), "La nouveauté de paradigme", font=get_font(31, True), fill="#16324F")
    d.text((165, 845), "L’addition ne subsiste que comme primitive cryptographique de SecAgg. L’objet appris par le serveur reste une frontière implicite,", font=get_font(27), fill="#1A1A1A")
    d.text((165, 892), "obtenue par résolution d’un problème géométrique non linéaire — et non une moyenne de modèles, gradients, poids ou logits.", font=get_font(27), fill="#1A1A1A")
    img.save(path, dpi=(180, 180))


def make_privacy_figure(path: Path):
    w, h = 2100, 940
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 40), "Trois garanties différentes — à ne pas confondre", font=get_font(44, True), fill="#16324F")
    cols = [
        ("Minimisation", "Sondes publiques\nPas d’images ni de points privés\nDescripteurs bornés", "#EAF2F8", "Réduit ce qui existe à protéger"),
        ("Agrégation sécurisée", "Masquage cryptographique\nSomme modulaire\nTolérance aux abandons", "#E7F3EF", "Cache les contributions individuelles"),
        ("DP distribuée", "Clipping\nParts de bruit discrètes\nComptabilité sur les rounds", "#FAEEDF", "Borne l’influence d’un client"),
    ]
    bw, gap, x0, y0, bh = 570, 95, 95, 205, 530
    for i, (title, body, fill, footer) in enumerate(cols):
        x = x0 + i*(bw+gap)
        rounded_box(d, (x, y0, x+bw, y0+bh), fill, outline="#2E74B5", radius=28, width=4)
        d.text((x+bw/2, y0+78), title, font=get_font(31, True), fill="#16324F", anchor="mm")
        yy = y0+165
        for line in body.split("\n"):
            d.text((x+52, yy), "• " + line, font=get_font(27), fill="#1A1A1A")
            yy += 62
        d.line((x+45, y0+405, x+bw-45, y0+405), fill="#A8B2BD", width=3)
        d.text((x+bw/2, y0+462), footer, font=get_font(25, True), fill="#5B6573", anchor="mm")
    draw_arrow(d, (540, 805), (1555, 805), fill="#2F7D68", width=7, head=20)
    d.text((1045, 855), "La protection crédible vient de l’empilement des trois couches", font=get_font(28, True), fill="#2F7D68", anchor="mm")
    img.save(path, dpi=(180, 180))


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Title", 30, NAVY, 0, 16),
        ("Subtitle", 15, GRAY, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = name in ("Title", "Heading 1", "Heading 2", "Heading 3")
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Subtitle"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Subtitle"].font.italic = True

    for list_name in ("List Bullet", "List Number"):
        st = styles[list_name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.194)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.208

    for name, parent, size, color, bold, italic in [
        ("Equation", "Normal", 10.5, NAVY, False, False),
        ("Caption Custom", "Normal", 9, GRAY, False, True),
        ("Callout Text", "Normal", 10.5, BLACK, False, False),
        ("Small Note", "Normal", 9, GRAY, False, False),
        ("Reference", "Normal", 9.2, BLACK, False, False),
    ]:
        if name not in styles:
            st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            st = styles[name]
        st.base_style = styles[parent]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = rgb(color)
        st.font.bold = bold
        st.font.italic = italic
        st.paragraph_format.space_after = Pt(5 if name != "Reference" else 3)
        if name == "Equation":
            st.font.name = "Cambria Math"
            st.paragraph_format.left_indent = Inches(0.25)
            st.paragraph_format.right_indent = Inches(0.25)
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def shade_paragraph(paragraph, fill=PALE, border=LIGHT_GRAY):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "5")
        e.set(qn("w:color"), border)
        p_bdr.append(e)
    p_pr.append(p_bdr)


def add_eq(doc, text, number=None):
    p = doc.add_paragraph(style="Equation")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text + (f"     ({number})" if number else ""))
    run.font.name = "Cambria Math"
    run.font.size = Pt(10.5)
    shade_paragraph(p, "F7F9FB", "D6DEE8")
    return p


def add_callout(doc, title, body, kind="idea"):
    cfg = {
        "idea": (LIGHT_BLUE, BLUE, "IDÉE CLÉ"),
        "warning": (LIGHT_ORANGE, ORANGE, "POINT DE VIGILANCE"),
        "privacy": (LIGHT_GREEN, GREEN, "CONFIDENTIALITÉ"),
        "claim": (LIGHT_RED, RED, "CLAIM DÉFENDABLE"),
    }
    fill, accent, default = cfg[kind]
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [170, 9190])
    left, right = table.rows[0].cells
    set_cell_shading(left, accent)
    set_cell_shading(right, fill)
    left.text = ""
    p = right.paragraphs[0]
    p.style = doc.styles["Callout Text"]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run((title or default) + "\n")
    r.bold = True
    r.font.color.rgb = rgb(accent)
    p.add_run(body)
    for c in (left, right):
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, PALE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = rgb(NAVY)
        r.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            r.font.size = Pt(font_size)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_table_width(table, widths)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.add_run(text)
    return p


def add_number(doc, text):
    # Une puce directionnelle évite qu'une numérotation Word continue
    # involontairement entre deux procédures éloignées du document.
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(text, style="Caption Custom")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    return p


def add_ref(doc, idx, citation, url):
    p = doc.add_paragraph(style="Reference")
    p.paragraph_format.left_indent = Inches(0.27)
    p.paragraph_format.first_line_indent = Inches(-0.27)
    p.add_run(f"[{idx}] {citation} ")
    add_hyperlink(p, "Accéder à la source", url)


def set_picture_alt(inline_shape, title: str, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_toc(doc):
    entries = [
        "1. Le changement de paradigme",
        "2. Cadre, notation et hypothèses",
        "3. Anatomie mathématique d’une frontière",
        "4. Champ de distance signée et SDF neuronale",
        "5. Reconstruction globale",
        "6. Confidentialité et DP distribuée",
        "7. Algorithme FedBCR-DP / FedEikonal",
        "8. Segmentation industrielle",
        "9. Applicabilité au-delà de la segmentation",
        "10. Positionnement bibliographique",
        "11. Programme théorique",
        "12. Parcours d’apprentissage",
        "13. Feuille de route expérimentale",
        "14. Risques et critères d’abandon",
        "15. Contribution proposée",
        "16. Conclusion et annexes",
    ]
    table = doc.add_table(rows=8, cols=2)
    set_table_width(table, [4680, 4680])
    for i, entry in enumerate(entries):
        row, col = i % 8, i // 8
        cell = table.cell(row, col)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(entry)
        r.font.size = Pt(10)
        r.font.color.rgb = rgb(NAVY)
        if entry.startswith(("1.", "6.", "10.", "12.")):
            r.bold = True
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tbl_pr.append(borders)


def add_cover(doc):
    section = doc.sections[0]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("NOTE DE RECHERCHE — CADRE CONCEPTUEL ET EXPÉRIMENTAL")
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = rgb(BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(66)
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("FedBCR-DP")
    r.font.name = "Calibri"
    r.font.size = Pt(38)
    r.font.bold = True
    r.font.color.rgb = rgb(NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("Apprendre une frontière globale sans moyenner les modèles")
    r.font.size = Pt(21)
    r.font.bold = True
    r.font.color.rgb = rgb(BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.55)
    p.paragraph_format.right_indent = Inches(0.55)
    r = p.add_run("Reconstruction fédérée de frontières de décision par champs de distance signée, contraintes d’Eikonal, transport optimal non équilibré et confidentialité différentielle distribuée")
    r.font.size = Pt(14)
    r.font.italic = True
    r.font.color.rgb = rgb(GRAY)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PARAMÈTRES  →  GÉOMÉTRIE  →  CHAMP GLOBAL")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = rgb(GREEN)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(96)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Document pédagogique de travail doctoral\nVersion du 17 juillet 2026")
    r.font.size = Pt(10.5)
    r.font.color.rgb = rgb(GRAY)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Statut scientifique : proposition à valider — aucune revendication de nouveauté absolue sans audit continu de la littérature et recherche d’antériorité.")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = rgb(RED)
    doc.add_page_break()


def configure_sections(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("FedBCR-DP  |  Cadre de recherche")
        r.font.size = Pt(8.5)
        r.font.color.rgb = rgb(GRAY)
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("—  ")
        r.font.size = Pt(8.5)
        r.font.color.rgb = rgb(GRAY)
        add_field(p.add_run(), "PAGE")
        r = p.add_run("  —")
        r.font.size = Pt(8.5)
        r.font.color.rgb = rgb(GRAY)
        first_header = section.first_page_header
        first_header.paragraphs[0].text = ""
        first_footer = section.first_page_footer
        first_footer.paragraphs[0].text = ""


def build_document():
    geometry = ASSET_DIR / "geometrie_certificat.png"
    pipeline = ASSET_DIR / "pipeline_fedbcr_dp.png"
    privacy = ASSET_DIR / "couches_confidentialite.png"
    make_geometry_figure(geometry)
    make_pipeline_figure(pipeline)
    make_privacy_figure(privacy)

    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    add_cover(doc)

    doc.add_heading("Résumé exécutif", level=1)
    doc.add_paragraph(
        "L’apprentissage fédéré classique construit un modèle global en agrégeant des paramètres ou des mises à jour. FedBCR — Federated Boundary-Constraint Reconstruction — propose de changer l’objet même de la collaboration : chaque client caractérise localement la géométrie de sa frontière de décision dans un espace latent commun; le serveur reconstruit ensuite une frontière globale comme le niveau zéro d’un champ de distance signée. FedEikonal désigne l’instanciation où ce champ est régularisé par l’équation d’Eikonal. La combinaison FedBCR-DP ajoute une esquisse additive, l’agrégation sécurisée et un bruit distribué afin que le serveur n’observe jamais les descripteurs individuels."
    )
    add_callout(
        doc,
        "La phrase qui résume la contribution",
        "Nous ne cherchons plus un barycentre dans l’espace des paramètres; nous cherchons une hypersurface globale compatible avec des contraintes géométriques locales, privées et hétérogènes.",
        "claim",
    )
    doc.add_paragraph(
        "Cette proposition est ambitieuse mais doit rester formulée avec rigueur. Les champs de distance signée, les réseaux implicites, l’équation d’Eikonal, le transport optimal, la confidentialité différentielle distribuée et la distillation de frontières existent séparément. La contribution potentielle réside dans leur assemblage précis au service d’une reconstruction fédérée de frontière — sous réserve qu’un audit bibliographique et une expérimentation reproductible confirment qu’aucune méthode antérieure n’effectue déjà cette combinaison."
    )
    add_table(doc, ["Question", "Réponse courte"], [
        ("Est-ce une variante de FedAvg ?", "Non au niveau de l’objet appris : aucune moyenne de poids, gradients ou logits. Une somme peut subsister comme primitive de SecAgg."),
        ("L’Eikonal impose-t-elle la segmentation ?", "Non. Elle impose une norme de gradient compatible avec une distance; elle s’applique à toute frontière dans un espace métrique pertinent."),
        ("Pourquoi commencer par la segmentation ?", "Parce que la frontière y est observable, mesurable et industriellement interprétable; c’est le banc d’essai le plus falsifiable."),
        ("Le certificat révèle-t-il le modèle ?", "Oui, potentiellement. D’où : sondes publiques, minimisation, clipping, esquisse additive, SecAgg et DP distribuée."),
        ("Le nom final ?", "FedBCR pour le cadre générique; FedEikonal pour la reconstruction SDF/Eikonal; FedBCR-DP pour l’architecture privée."),
    ], [1900, 7460])

    doc.add_heading("Table des matières", level=1)
    add_toc(doc)
    p = doc.add_paragraph(style="Small Note")
    p.add_run("Plan de lecture matérialisé pour rester visible dans tous les moteurs Word/PDF.")
    doc.add_page_break()

    doc.add_heading("1. Le changement de paradigme", level=1)
    doc.add_heading("1.1 Ce que fait FedAvg", level=2)
    doc.add_paragraph(
        "Au round t, le serveur envoie un modèle θᵗ. Le client k l’entraîne sur ses données locales Dₖ et renvoie θₖᵗ⁺¹ ou une mise à jour Δθₖᵗ. FedAvg calcule un barycentre pondéré par la quantité de données nₖ [R1] :"
    )
    add_eq(doc, "θᵗ⁺¹ = Σₖ₌₁ᴷ (nₖ / Σⱼ nⱼ) θₖᵗ⁺¹.", "1")
    doc.add_paragraph(
        "Ce calcul est simple et efficace lorsque les modèles partagent exactement la même paramétrisation. Mais la moyenne se fait dans l’espace des paramètres, où deux réseaux fonctionnellement proches peuvent être éloignés à cause des permutations de neurones, et où l’hétérogénéité non-IID peut faire diverger les optima locaux. Les nombreuses alternatives à FedAvg changent les pondérations, ajoutent des termes proximaux, distillent des prédictions ou alignent des représentations; elles restent souvent dans une logique d’agrégation d’objets statistiques ou paramétriques [R2]."
    )
    doc.add_heading("1.2 L’objet de FedBCR", level=2)
    doc.add_paragraph(
        "FedBCR part d’une autre question : quelles régions de l’espace représentent les classes, et où les décisions changent-elles ? Le serveur ne cherche pas d’abord un vecteur θ global. Il cherche une hypersurface Γ⋆, puis un champ scalaire φ⋆ dont le niveau zéro représente cette hypersurface :"
    )
    add_eq(doc, "Γ⋆ = {z ∈ ℝᵈ : φ⋆(z)=0},       signe(φ⋆(z)) = classe prédite.", "2")
    doc.add_paragraph(
        "Cette formulation est fonctionnelle et géométrique. Elle accepte, en principe, des têtes locales de tailles ou d’architectures différentes, à condition qu’elles évaluent leurs décisions dans un repère latent commun. Le résultat global peut être utilisé directement comme classifieur implicite, distillé dans un réseau compact, ou servir de régularisateur à une étape locale suivante."
    )
    shape = doc.add_picture(str(pipeline), width=Inches(6.5))
    set_picture_alt(shape, "Pipeline FedBCR-DP", "Six étapes : repère latent commun, géométrie locale, esquisse privée, agrégation sécurisée, reconstruction SDF/Eikonal et modèle global utilisable.")
    add_caption(doc, "Figure 1 — Chaîne conceptuelle FedBCR-DP. La non-linéarité principale se situe dans la reconstruction géométrique au serveur.")
    add_callout(doc, "Ce qui est — et n’est pas — nouveau", "La nouveauté ne peut pas être “utiliser une SDF”, “utiliser Eikonal” ou “utiliser la DP”. Elle doit être revendiquée comme une architecture et un problème d’optimisation précis : reconstruction fédérée privée d’une frontière de décision depuis des contraintes locales orientées, sans fusion des modèles.", "warning")

    doc.add_heading("2. Cadre, notation et hypothèses", level=1)
    add_table(doc, ["Symbole", "Définition"], [
        ("K, k", "nombre de clients et indice d’un client"),
        ("Dₖ", "jeu de données privé du client k"),
        ("x, y", "observation et étiquette; y ∈ {1,…,C}"),
        ("E", "encodeur partagé ou publiquement spécifié, E : 𝒳 → ℝᵈ"),
        ("z=E(x)", "représentation latente commune"),
        ("hₖ", "tête de décision locale du client k"),
        ("sₖ,𝑐(z)", "score réel — logit avant softmax — attribué à la classe c par hₖ"),
        ("mₖ,ₐ,ᵦ(z)", "marge paire-à-paire entre les classes a et b"),
        ("Γₖ,ₐᵦ", "frontière locale active entre a et b"),
        ("z_q", "sonde publique ou synthétique, identique pour tous les clients"),
        ("nₖ,q", "normale unitaire locale estimée près de z_q"),
        ("δₖ,q", "déplacement signé de z_q vers la frontière locale"),
        ("φψ", "champ de distance signée neuronal, paramétré par ψ"),
        ("μΓ", "mesure portée par une frontière Γ"),
    ], [1450, 7910])
    doc.add_heading("2.1 Le repère commun est une hypothèse structurante", level=2)
    doc.add_paragraph(
        "Comparer des frontières n’a de sens que si les coordonnées latentes ont la même signification chez tous les clients. Un encodeur E entraîné indépendamment sur chaque client ne fournit pas automatiquement ce repère : les espaces peuvent être tournés, permutés ou déformés. La première version de FedBCR doit donc utiliser l’une des options suivantes."
    )
    add_bullet(doc, "Encodeur public pré-entraîné et gelé — option la plus propre pour la preuve de concept.")
    add_bullet(doc, "Encodeur partagé, entraîné en amont puis gelé pendant la reconstruction de frontière.")
    add_bullet(doc, "Alignement explicite des espaces latents par sondes publiques appariées — extension plus difficile.")
    add_callout(doc, "Hypothèse H0", "Tous les clients évaluent leurs têtes dans le même espace latent métrique (ℝᵈ, ‖·‖₂). Sans H0, la normale et la distance de deux clients ne sont pas comparables.", "idea")

    doc.add_heading("3. Anatomie mathématique d’une frontière de décision", level=1)
    doc.add_heading("3.1 Que représente sₖ,𝑐(z) ?", level=2)
    doc.add_paragraph(
        "La tête locale hₖ prend un vecteur latent z et renvoie C scores : hₖ(z)=(sₖ,₁(z),…,sₖ,C(z)). Le score sₖ,𝑐(z) est généralement le logit de la classe c, c’est-à-dire la quantité réelle avant la transformation softmax. La classe prédite est :"
    )
    add_eq(doc, "ŷₖ(z) = arg max_{c∈{1,…,C}} sₖ,c(z).", "3")
    doc.add_paragraph(
        "Pour étudier le passage de la classe a à la classe b, on définit la marge paire-à-paire :"
    )
    add_eq(doc, "mₖ,ₐ,ᵦ(z) = sₖ,ₐ(z) − sₖ,ᵦ(z).", "4")
    doc.add_paragraph(
        "Le signe de m indique quel score domine entre a et b. Dans le cas binaire, m=0 décrit directement la frontière. Dans le cas multiclasse, m=0 contient aussi des points où une troisième classe c domine. La frontière active correcte est donc :"
    )
    add_eq(doc, "Γₖ,ₐᵦ = {z : sₖ,ₐ(z)=sₖ,ᵦ(z)=max_c sₖ,c(z)}.", "5")
    add_callout(doc, "Pourquoi cette précision compte", "Utiliser seulement {mₖ,ₐᵦ=0} en multiclasse enverrait des contraintes appartenant à des égalités non décisionnelles. Le filtre d’activité évite de reconstruire des morceaux de surface qui ne changent aucune prédiction.", "warning")

    doc.add_heading("3.2 D’où vient la normale nₖ,q ?", level=2)
    doc.add_paragraph(
        "Supposons que m soit différentiable et que ∇m(z)≠0 sur la frontière. Le théorème des fonctions implicites dit que {m=0} est localement une hypersurface régulière. Si v est un vecteur tangent à cette surface, une courbe γ(t) contenue dans la surface vérifie m(γ(t))=0. En dérivant à t=0 :"
    )
    add_eq(doc, "d/dt m(γ(t))|₀ = ∇m(z)ᵀ γ′(0) = ∇m(z)ᵀv = 0.", "6")
    doc.add_paragraph(
        "Le gradient est donc orthogonal à tous les vecteurs tangents : il donne la direction normale. Sa normalisation définit :"
    )
    add_eq(doc, "nₖ,q = ∇mₖ,ₐ,ᵦ(z_q) / (‖∇mₖ,ₐ,ᵦ(z_q)‖₂ + ηₙ).", "7")
    doc.add_paragraph(
        "ηₙ>0 évite une division instable mais introduit un léger biais. En pratique, on rejette une sonde lorsque ‖∇m‖ est inférieur à un seuil τ_grad; une normale produite dans une zone plate n’est pas fiable. L’orientation doit aussi être conventionnelle, par exemple du côté b vers le côté a, puisque ∇(sₐ−sᵦ) pointe vers l’augmentation relative de a."
    )

    doc.add_heading("3.3 D’où vient le déplacement δₖ,q ?", level=2)
    doc.add_paragraph(
        "On veut déplacer la sonde z_q le long de la normale n jusqu’à un point b=z_q+δn où la marge s’annule. Le développement de Taylor au premier ordre donne :"
    )
    add_eq(doc, "m(z_q+δn) ≈ m(z_q) + δ ∇m(z_q)ᵀn.", "8")
    doc.add_paragraph(
        "Avec n=∇m/‖∇m‖, le produit scalaire vaut ‖∇m‖. En imposant l’approximation m(z_q+δn)=0 :"
    )
    add_eq(doc, "δₖ,q ≈ − mₖ,ₐ,ᵦ(z_q) / (‖∇mₖ,ₐ,ᵦ(z_q)‖₂ + ηδ).", "9")
    doc.add_paragraph(
        "Cette formule n’est pas une distance exacte; c’est une estimation linéaire locale. Elle est fiable lorsque z_q est suffisamment proche de la frontière et que la courbure de m reste modérée dans le voisinage. Une projection itérative de type Newton normalisé est plus précise :"
    )
    add_eq(doc, "z⁽ʳ⁺¹⁾ = z⁽ʳ⁾ − [m(z⁽ʳ⁾)/(‖∇m(z⁽ʳ⁾)‖₂²+η)] ∇m(z⁽ʳ⁾).", "10")
    doc.add_paragraph(
        "On arrête lorsque |m(z⁽ʳ⁾)|≤τ_m, lorsque le déplacement devient petit, ou après R_max itérations. Le point projeté bₖ,q, la normale recalculée en bₖ,q et un résidu final forment un descripteur plus fidèle que la seule formule au premier ordre."
    )
    shape = doc.add_picture(str(geometry), width=Inches(6.5))
    set_picture_alt(shape, "Descripteur local orienté de frontière", "Une sonde publique est projetée vers la frontière locale; le déplacement signé, la normale, la paire de classes et la confiance constituent le descripteur.")
    add_caption(doc, "Figure 2 — Construction locale d’un descripteur orienté. Le terme “certificat” désigne ici une attestation numérique de géométrie, pas une preuve cryptographique.")

    doc.add_heading("3.4 Le “certificat géométrique” décortiqué", level=2)
    doc.add_paragraph(
        "Pour éviter une promesse excessive, le terme recommandé dans l’article est descripteur local orienté de frontière. Un certificat Cₖ,q peut être représenté par :"
    )
    add_eq(doc, "Cₖ,q = (q, a, b, δₖ,q, nₖ,q, rₖ,q, ρₖ,q, wₖ,q).", "11")
    add_table(doc, ["Champ", "Sens", "Contrôle de qualité"], [
        ("q", "identifiant d’une sonde publique", "aucune coordonnée privée transmise"),
        ("a,b", "classes séparées", "égalité active au sens de (5)"),
        ("δ", "décalage signé", "clip |δ|≤δ_max"),
        ("n", "orientation locale", "‖n‖≈1 et convention de signe"),
        ("r", "résidu |m(b)|", "rejeter si r>τ_m"),
        ("ρ", "rayon local de validité", "dérivé de la courbure ou d’un test de stabilité"),
        ("w", "poids de confiance", "fonction bornée de r, ‖∇m‖, stabilité et support local"),
    ], [900, 4080, 4380])
    add_callout(doc, "Le descripteur n’est pas une vérité", "Il résume une approximation locale du comportement du modèle k. Il doit être accompagné d’un résidu, d’un seuil d’activité et d’une estimation de validité. Sans ces contrôles, le serveur fusionne des artefacts numériques comme s’ils étaient de la géométrie.", "warning")

    doc.add_heading("4. Champ de distance signée et SDF neuronale", level=1)
    doc.add_heading("4.1 SDF classique", level=2)
    doc.add_paragraph(
        "Soit Ω une région de classe positive et Γ=∂Ω sa frontière. La distance signée associe à chaque z la distance euclidienne à Γ, avec un signe indiquant le côté :"
    )
    add_eq(doc, "d_Γ(z) = +dist(z,Γ) si z∈Ω;  0 si z∈Γ;  −dist(z,Γ) sinon.", "12")
    doc.add_paragraph(
        "Loin de l’axe médian et des points de non-différentiabilité, la distance signée satisfait l’équation d’Eikonal :"
    )
    add_eq(doc, "‖∇d_Γ(z)‖₂ = 1 presque partout, avec d_Γ(z)=0 sur Γ.", "13")
    doc.add_paragraph(
        "Cette équation dit que le champ varie à vitesse unité lorsque l’on s’éloigne normalement de la frontière. Elle ne crée pas la frontière à elle seule : il faut des conditions au bord, des signes ou des observations. Elle n’est pas non plus valide de manière lisse partout; la projection sur Γ peut être non unique sur l’axe médian."
    )

    doc.add_heading("4.2 Qu’est-ce qu’une SDF neuronale ?", level=2)
    doc.add_paragraph(
        "Une SDF neuronale est une fonction φψ:ℝᵈ→ℝ représentée par un réseau — souvent un perceptron multicouche — qui approxime un champ de distance signée continu. Sa frontière implicite est {φψ=0}. Cette famille appartient aux représentations neuronales implicites : au lieu de stocker une grille ou un maillage, on interroge le réseau en n’importe quel point [R3–R4]."
    )
    doc.add_paragraph(
        "Dans FedBCR, les descripteurs locaux fournissent des contraintes sur la valeur zéro, la normale et le signe. Une perte type est :"
    )
    add_eq(doc, "L(ψ)=λ₀L₀ + λ_Eik L_Eik + λ_nL_n + λ_sL_s + λ_regL_reg.", "14")
    add_table(doc, ["Terme", "Formule indicative", "Rôle"], [
        ("L₀", "E_b |φψ(b)|", "place les points reconstruits sur le niveau zéro"),
        ("L_Eik", "E_z (‖∇φψ(z)‖₂−1)²", "rend le champ compatible avec une distance"),
        ("L_n", "E_(b,n)[1−⟨∇φ/‖∇φ‖,n⟩]", "aligne la normale globale avec les orientations locales"),
        ("L_s", "E_(z,y) softplus(−yφψ(z))", "fixe le côté positif/négatif"),
        ("L_reg", "courbure, lissage ou parcimonie", "évite les oscillations non justifiées"),
    ], [1100, 3620, 4640], 9.1)
    doc.add_paragraph(
        "Les points b ne sont pas nécessairement transmis individuellement dans la version privée; ces pertes peuvent être estimées à partir d’une esquisse agrégée sur une base commune. Le terme Eikonal est un régularisateur, pas une garantie que φψ est l’unique vraie distance. Les poids λ doivent être étudiés par ablation et la qualité doit être jugée sur la frontière, la calibration et la tâche finale."
    )

    doc.add_heading("4.3 Binaire, multiclasse et segmentation", level=2)
    add_bullet(doc, "Binaire : un seul champ φ; signe(φ) donne les deux régions.")
    add_bullet(doc, "Multiclasse paire-à-paire : un champ φₐᵦ par paire active; nécessite une règle de cohérence pour éviter les cycles a>b, b>c, c>a.")
    add_bullet(doc, "Multiclasse par potentiels : apprendre C potentiels g_c(z), puis prédire argmax_c g_c; les frontières sont g_a=g_b. Plus cohérent, mais moins directement SDF.")
    add_bullet(doc, "Segmentation sémantique : φψ(I,u,c) dépend de l’image I, de la coordonnée pixel/voxel u et éventuellement de la classe c. Le niveau zéro trace un contour conditionné par l’image.")
    add_bullet(doc, "Segmentation d’instances : il faut en plus distinguer l’identité des objets; un unique champ par classe ne suffit pas. C’est une extension de seconde génération.")
    add_callout(doc, "Réponse directe", "FedEikonal n’est pas réservé à la segmentation. La classification simple est possible dans un espace latent partagé. La segmentation est cependant le premier terrain recommandé, car la frontière a une signification spatiale observable et les métriques industrielles sont parlantes.", "idea")

    doc.add_heading("5. Reconstruction globale : de contraintes locales à une frontière", level=1)
    doc.add_heading("5.1 Pourquoi Hausdorff seule ne suffit pas", level=2)
    doc.add_paragraph(
        "La distance de Hausdorff symétrique entre deux ensembles A et B est :"
    )
    add_eq(doc, "d_H(A,B)=max{sup_{a∈A} inf_{b∈B}‖a−b‖, sup_{b∈B} inf_{a∈A}‖b−a‖}.", "15")
    doc.add_paragraph(
        "Elle mesure le pire écart spatial. Elle est utile pour un défaut local critique, mais très sensible aux valeurs aberrantes et ne pénalise pas directement une différence de longueur, de masse, de multiplicité ou de topologie. Deux contours de longueurs très différentes peuvent être géométriquement proches au sens de Hausdorff si chacun reste dans un tube étroit autour de l’autre."
    )
    doc.add_paragraph(
        "La distance de Hausdorff modifiée de Dubuisson–Jain remplace l’extrême par une moyenne des plus proches voisins [R5]. Elle est plus robuste aux points aberrants, mais peut masquer un petit segment industriel critique : une fissure courte mais dangereuse contribue peu à une moyenne globale. Elle doit donc être une métrique de diagnostic, pas l’unique objectif."
    )

    doc.add_heading("5.2 Représenter la longueur et la masse", level=2)
    doc.add_paragraph(
        "On associe à une frontière Γ une mesure de surface μ_Γ=ℋ^{d−1}|_Γ, où ℋ^{d−1} est la mesure de Hausdorff de dimension d−1. En 2D, sa masse totale est la longueur du contour; en 3D, son aire. Le transport optimal non équilibré — UOT — compare la position et autorise la création/destruction de masse avec un coût explicite [R6–R7]. Il peut donc pénaliser un contour trop court ou trop long, contrairement à une simple proximité d’ensembles."
    )
    add_eq(doc, "UOT(μ,ν)=inf_π  ⟨C,π⟩ + τ₁D(π₁‖μ) + τ₂D(π₂‖ν).", "16")
    doc.add_paragraph(
        "π est un plan de transport, C le coût de déplacement, π₁ et π₂ ses marginales, et D une divergence — souvent de type Kullback–Leibler — qui facture le défaut de masse. Pour l’industrie, on recommande une hiérarchie plutôt qu’une somme arbitraire de métriques :"
    )
    add_number(doc, "Respecter d’abord les contraintes topologiques essentielles : nombre de composantes, absence de trou impossible, connexité d’une fissure.")
    add_number(doc, "Minimiser ensuite une divergence UOT sur les mesures de contour, éventuellement pondérée par zones critiques.")
    add_number(doc, "Contrôler enfin HD95 ou le Hausdorff dirigé dans les régions où le pire écart est physiquement important.")
    add_number(doc, "Reporter MHD, longueur relative, courbure et métriques de segmentation comme diagnostics complémentaires.")

    doc.add_heading("5.3 Objectif de reconstruction proposé", level=2)
    doc.add_paragraph(
        "Une formulation de recherche, à préciser expérimentalement, est :"
    )
    add_eq(doc, "min_ψ  L_data(φψ; Ũ) + λ_EikL_Eik + λ_topL_top + λ_curvL_curv,", "17")
    add_eq(doc, "sous contraintes : cohérence de signe, activité multiclasse et budget de complexité.", "18")
    doc.add_paragraph(
        "Ũ désigne l’esquisse agrégée et bruitée. L_data peut mesurer la compatibilité avec des moments de points de frontière, des orientations et une masse par cellule, ou une divergence UOT reconstruite depuis ces statistiques. La topologie peut être évaluée via nombres de Betti ou persistance topologique; cette composante est optionnelle pour un premier prototype, car elle augmente fortement la complexité."
    )

    doc.add_heading("6. Le problème de confidentialité", level=1)
    doc.add_heading("6.1 Pourquoi une frontière peut fuiter", level=2)
    doc.add_paragraph(
        "Une frontière de décision révèle les régions où le modèle hésite, ses directions sensibles et parfois la structure du support de données. Une collection dense de points exacts et de normales peut faciliter la copie du modèle, la génération d’exemples adversariaux, l’inférence de propriétés ou la reconstruction approximative du comportement local [R8]. Le fait de ne pas transmettre les images ne suffit donc pas à affirmer la confidentialité."
    )
    add_callout(doc, "Principe", "Un descripteur géométrique n’est pas anonyme par nature. Plus il est exact, dense et individualisé, plus il peut devenir une empreinte du client.", "privacy")
    shape = doc.add_picture(str(privacy), width=Inches(6.5))
    set_picture_alt(shape, "Couches de confidentialité", "La minimisation réduit les informations, l’agrégation sécurisée cache les contributions individuelles et la DP distribuée borne l’influence d’un client.")
    add_caption(doc, "Figure 3 — Minimisation, agrégation sécurisée et DP répondent à trois questions différentes.")

    doc.add_heading("6.2 Définition de la confidentialité différentielle", level=2)
    doc.add_paragraph(
        "Un mécanisme aléatoire M est (ε,δ)-différentiellement privé si, pour deux jeux voisins D et D′ et tout ensemble de sorties S :"
    )
    add_eq(doc, "Pr[M(D)∈S] ≤ e^ε Pr[M(D′)∈S] + δ.", "19")
    doc.add_paragraph(
        "Le sens de voisin doit être annoncé. En DP au niveau client, D et D′ diffèrent par la participation entière d’un client. En DP au niveau image/patient, ils diffèrent par un seul enregistrement. La première protège une organisation ou un appareil; la seconde protège un individu au sein d’un silo. Les sensibilités, les mécanismes et l’utilité ne sont pas identiques [R9–R10]."
    )
    add_table(doc, ["Régime", "Ce que voit le serveur", "Confiance", "Conséquence"], [
        ("DP centrale", "agrégat bruité", "serveur ou composant de confiance avant bruit", "meilleure utilité en général"),
        ("DP locale", "messages déjà privés individuellement", "aucune confiance", "bruit souvent élevé"),
        ("DP distribuée + SecAgg", "somme bruitée des messages", "seuil d’honnêteté et protocole cryptographique", "utilité proche de la DP centrale possible"),
    ], [1500, 2700, 2400, 2760], 9.0)

    doc.add_heading("6.3 L’incompatibilité cachée avec SecAgg", level=2)
    doc.add_paragraph(
        "Les protocoles d’agrégation sécurisée classiques révèlent une somme modulaire et rien d’autre [R11]. Or UOT, l’appariement de points, les contraintes topologiques et l’optimisation d’une SDF sont non linéaires. Le serveur ne peut pas calculer directement ces opérations sur des certificats individuels qu’il ne voit pas. Deux solutions existent."
    )
    add_bullet(doc, "Solution pratique recommandée : chaque client encode sa géométrie dans une esquisse vectorielle additive définie sur des sondes ou une base commune. SecAgg additionne les esquisses; le serveur effectue ensuite la reconstruction non linéaire sur l’agrégat.")
    add_bullet(doc, "Solution cryptographique lourde : calcul multipartite sécurisé ou chiffrement homomorphe pour évaluer UOT et la reconstruction sur des certificats chiffrés. Plus expressive, mais coûteuse et difficile à industrialiser.")
    add_callout(doc, "Une somme réapparaît — sans annuler le paradigme", "La somme de SecAgg est un substrat cryptographique/statistique. Nous ne faisons toujours pas une moyenne de modèles, gradients, paramètres ou logits. Le serveur résout ensuite un problème géométrique non linéaire pour construire φ⋆.", "idea")

    doc.add_heading("6.4 Esquisse additive privée", level=2)
    doc.add_paragraph(
        "Soit Φ(b,n,a,b)∈ℝᴰ une base commune : cellules spatiales, fonctions radiales RBF, ondelettes, random features et bins d’orientation sont des candidats. Le client transforme ses Jₖ descripteurs en :"
    )
    add_eq(doc, "uₖ = Σⱼ₌₁ᴶᵏ wₖ,j Φ(bₖ,j,nₖ,j,aⱼ,bⱼ).", "20")
    doc.add_paragraph("Il borne ensuite la contribution complète du client :")
    add_eq(doc, "ūₖ = uₖ · min(1, C/‖uₖ‖₂).", "21")
    doc.add_paragraph(
        "Sous une adjacence ajout/retrait, la sensibilité L₂ de la somme est au plus C; sous une adjacence remplacement, elle est au plus 2C. Après quantification Q_s, chaque client ajoute une part de bruit ξₖ et participe à SecAgg :"
    )
    add_eq(doc, "Ũ = Decode([Σₖ (Q_s(ūₖ)+ξₖ)] mod M).", "22")
    doc.add_paragraph(
        "Pour une approximation gaussienne, si au moins h_min clients honnêtes survivent au protocole et si la variance agrégée cible est σ²Δ² par coordonnée, chaque client honnête peut contribuer une variance τ²=σ²Δ²/h_min. En pratique, le bruit doit être compatible avec l’arithmétique discrète et modulaire : gaussienne discrète, Skellam ou mécanisme binomial de Poisson [R12–R14]. La calibration exacte doit suivre le théorème du mécanisme choisi, pas la seule intuition gaussienne."
    )

    doc.add_heading("6.5 Menaces, collusion, abandons et composition", level=2)
    add_table(doc, ["Menace", "Réponse", "Ce qui reste ouvert"], [
        ("Serveur curieux", "SecAgg + sortie DP", "métadonnées, taille des cohortes, sondes adaptatives"),
        ("Clients collusifs", "calibrer pour h_min honnêtes", "estimer le seuil réaliste par déploiement"),
        ("Abandons", "protocole SecAgg tolérant aux dropouts", "bruit résiduel suffisant après abandon"),
        ("Client malveillant", "clipping vérifiable, règles robustes", "DP ne garantit pas l’intégrité"),
        ("Répétition des rounds", "comptable RDP/PLD", "budget ε total et arrêt de confidentialité"),
        ("Sondes adaptatives", "pré-enregistrement ou audit", "risque d’interrogation ciblée"),
    ], [1650, 3170, 4540], 9.0)
    doc.add_paragraph(
        "La composition est cruciale : publier T agrégats privés consomme davantage de budget qu’un seul. Une étude sérieuse fixe δ avant l’entraînement, utilise un comptable de confidentialité par round — par exemple RDP — et arrête le protocole lorsque ε atteint le budget annoncé. Dans un petit scénario cross-silo, le nombre K de clients peut être trop faible pour amortir le bruit; il faut alors préférer des statistiques plus basses dimensions, réduire la fréquence des reconstructions, ou viser la DP au niveau image calculée à l’intérieur de chaque silo."
    )

    doc.add_heading("7. Algorithme FedBCR-DP / FedEikonal", level=1)
    doc.add_heading("7.1 Phase de préparation", level=2)
    add_number(doc, "Choisir l’espace commun : encodeur E public et gelé, dimension latente d contrôlée.")
    add_number(doc, "Construire une banque de sondes Z_pub={z_q} depuis des données publiques, synthétiques ou générées sous contraintes.")
    add_number(doc, "Définir la base additive Φ, la norme de clipping C, le module M, la quantification s, h_min et le budget (ε,δ).")
    add_number(doc, "Pré-enregistrer les seuils τ_grad, τ_m, δ_max, l’activité multiclasse et le nombre maximal de descripteurs par client.")

    doc.add_heading("7.2 Round client k", level=2)
    add_number(doc, "Entraîner ou mettre à jour la tête locale hₖ sur Dₖ, sans transmettre ses paramètres.")
    add_number(doc, "Pour chaque sonde et paire candidate, calculer les logits sₖ,c(z_q), la marge m et ∇m par différentiation automatique.")
    add_number(doc, "Vérifier l’activité, projeter vers m=0 avec (10), recalculer normale, résidu et confiance.")
    add_number(doc, "Encoder les descripteurs dans uₖ avec (20), puis clipper selon (21).")
    add_number(doc, "Quantifier, ajouter la part de bruit distribuée et envoyer le message masqué au protocole SecAgg.")

    doc.add_heading("7.3 Round serveur", level=2)
    add_number(doc, "Récupérer uniquement l’esquisse agrégée bruitée Ũ.")
    add_number(doc, "Détecter les cellules, orientations et paires de classes suffisamment soutenues.")
    add_number(doc, "Optimiser φψ avec la fidélité à Ũ, Eikonal, orientation, signe et régularisation.")
    add_number(doc, "Valider la géométrie sur un jeu public; appliquer des critères d’arrêt et un audit du budget DP.")
    add_number(doc, "Publier φψ, ou distiller ses décisions dans un modèle compact; ne jamais publier les contributions individuelles.")
    add_callout(doc, "Invariant architectural", "Le serveur ne reçoit aucune tête locale hₖ, aucun gradient d’entraînement, aucune image privée et aucun point de frontière individualisé dans la variante DP recommandée.", "privacy")

    doc.add_heading("8. Pourquoi l’expérimentation initiale doit être une segmentation industrielle", level=1)
    doc.add_heading("8.1 Le choix recommandé", level=2)
    doc.add_paragraph(
        "Commencer par une segmentation sémantique binaire de défauts — fond versus défaut — sur images industrielles. Ce choix ne réduit pas l’ambition; il rend l’hypothèse falsifiable. La frontière est visible, la distance signée peut être calculée depuis les masques de vérité terrain, les défauts fins mettent immédiatement en évidence les limites de Hausdorff/MHD, et la longueur, la connectivité et les erreurs locales ont une signification industrielle."
    )
    doc.add_paragraph(
        "Pour une image I et un pixel u∈Ω_img, le modèle local fournit des logits sₖ,c(I,u). La marge binaire est mₖ(I,u)=sₖ,defaut(I,u)−sₖ,fond(I,u). La SDF globale doit être conditionnelle à l’image :"
    )
    add_eq(doc, "φψ(I,u) ≈ distance signée du pixel u à la frontière du défaut dans I.", "23")
    doc.add_paragraph(
        "Il ne faut pas confondre deux espaces. Pour la segmentation spatiale, ∇_uφ décrit la normale du contour dans le plan image. Pour une frontière de classification latente, ∇_z m décrit une direction dans l’espace des caractéristiques. Le prototype peut étudier d’abord la géométrie spatiale, puis relier les deux en faisant dépendre φ de caractéristiques E(I,u)."
    )

    doc.add_heading("8.2 Datasets et partition fédérée", level=2)
    add_table(doc, ["Élément", "Recommandation"], [
        ("Jeux de données", "MVTec AD/LOCO pour anomalies, KolektorSDD/SDD2 pour défauts fins, DAGM ou Severstal selon licences et masques disponibles."),
        ("Clients", "sites, machines, lignes ou types de texture; éviter une partition aléatoire qui gomme l’hétérogénéité."),
        ("Non-IID", "Dirichlet sur fréquence des défauts + décalage de capteur/texture + déséquilibre de taille."),
        ("Sondes", "images publiques/synthétiques et coordonnées; aucune image locale utilisée comme sonde serveur."),
        ("Première tâche", "binaire, 2D, contours fermés ou défauts filiformes; 5 à 20 clients simulés."),
        ("Validation externe", "un site jamais vu et un capteur décalé pour tester la reconstruction hors distribution."),
    ], [1700, 7660], 9.2)

    doc.add_heading("8.3 Baselines", level=2)
    add_bullet(doc, "Local-only et centralisé — bornes basse et haute, sans les confondre avec des méthodes fédérées.")
    add_bullet(doc, "FedAvg, FedProx et une méthode de distillation fédérée adaptée à la segmentation.")
    add_bullet(doc, "Agrégation de probabilités/logits sur les mêmes sondes publiques.")
    add_bullet(doc, "FedBCR sans Eikonal, sans normale, sans UOT, puis version complète.")
    add_bullet(doc, "FedBCR-DP à budgets croissants et SecAgg seul — pour séparer confidentialité cryptographique et statistique.")
    add_bullet(doc, "Oracle de reconstruction utilisant des contours non privés agrégés — seulement pour mesurer le coût de confidentialité.")

    doc.add_heading("8.4 Métriques : ne pas laisser Dice cacher la frontière", level=2)
    add_table(doc, ["Famille", "Métriques"], [
        ("Région", "Dice, IoU, précision/rappel défaut, AUROC si pertinent"),
        ("Frontière", "Boundary F-score à plusieurs tolérances, ASSD, HD95, Hausdorff dirigé en zones critiques"),
        ("Masse/forme", "erreur relative de longueur, UOT de contour, distribution de courbure"),
        ("Topologie", "nombre de composantes, erreurs de connexité, Betti-0/Betti-1 si justifié"),
        ("Décision", "calibration ECE/Brier, robustesse OOD, marge"),
        ("Fédération", "octets montants/descendants, nombre de rounds, temps mur, variance inter-clients"),
        ("Énergie", "joules compute/communication, énergie jusqu’à une cible de Dice, survie des clients selon le cadre du dépôt"),
        ("Confidentialité", "ε, δ, unité protégée, h_min, taux d’abandon, avantage d’attaque"),
    ], [1600, 7760], 9.1)

    doc.add_heading("8.5 Plan d’ablation minimal", level=2)
    add_table(doc, ["Ablation", "Question scientifique"], [
        ("− normale", "la simple localisation des frontières suffit-elle ?"),
        ("− Eikonal", "le champ apprend-il une surface mais une mauvaise distance ?"),
        ("MHD ↔ UOT", "la masse/longueur améliore-t-elle les défauts fins ?"),
        ("points ↔ esquisse", "quel coût vient de la compression additive ?"),
        ("sans DP ↔ ε∈{1,2,4,8}", "courbe confidentialité–géométrie–tâche"),
        ("d et nombre de sondes", "où se situe le compromis dimension–communication ?"),
        ("E gelé ↔ aligné", "quelle dépendance au repère commun ?"),
        ("K, non-IID, dropouts", "stabilité fédérée et validité de h_min"),
    ], [2600, 6760], 9.2)

    doc.add_heading("9. Applicabilité au-delà de la segmentation", level=1)
    add_table(doc, ["Tâche", "Applicabilité", "Adaptation"], [
        ("Classification tabulaire/image", "forte", "SDF dans un latent commun; sondes publiques ou génératives"),
        ("Segmentation sémantique", "très forte", "SDF conditionnelle à l’image et au pixel/voxel"),
        ("Segmentation d’instances", "moyenne", "champs multi-instance, embeddings d’objets ou level sets multiples"),
        ("NLP classification", "possible mais délicate", "frontière dans l’espace d’un encodeur gelé; la distance latente doit être justifiée"),
        ("Séquence/génération", "faible en l’état", "pas de frontière unique; nécessiterait contraintes sur trajectoires ou distributions"),
        ("Régression", "indirecte", "reconstruire des level sets f(z)=τ pour plusieurs seuils, donc une autre formulation"),
        ("Graphes", "possible", "p-Eikonal/Hamilton–Jacobi sur graphe, sans grille euclidienne [R15]"),
    ], [1850, 1500, 6010], 9.0)
    doc.add_paragraph(
        "En haute dimension, résoudre Eikonal sur une grille régulière est impossible : le nombre de cellules croît exponentiellement avec d. Il faut une SDF neuronale, un ensemble adaptatif de sondes, une base creuse, ou une formulation sur graphe. En NLP, la norme euclidienne du latent peut ne pas correspondre à une variation sémantique; utiliser Eikonal sans valider la métrique serait une erreur conceptuelle."
    )

    doc.add_heading("10. Positionnement bibliographique et revendication de nouveauté", level=1)
    doc.add_heading("10.1 Familles voisines", level=2)
    add_table(doc, ["Famille", "Proximité", "Différence proposée"], [
        ("FedAvg et variantes", "modèle global fédéré", "agrègent paramètres/mises à jour; FedBCR reconstruit une surface"),
        ("Distillation fédérée", "partage de sorties ou connaissances", "FedBCR partage des contraintes différentielles de frontière, pas seulement des logits"),
        ("Fed-GDBD", "distillation orientée frontière globale", "emploie prototypes et distillation; ne reconstruit pas explicitement une SDF Eikonal privée [R16]"),
        ("DeepSDF / IGR", "surface implicite et Eikonal", "non fédéré, principalement reconstruction géométrique 3D [R3–R4]"),
        ("Boundary loss / HD loss", "segmentation guidée par frontière", "pertes locales/centralisées, pas une agrégation fédérée de frontière [R17–R18]"),
        ("p-Eikonal sur graphes", "classification via HJ/Eikonal", "non fédéré; exploite un graphe de données [R15]"),
        ("DP distribuée + SecAgg", "bruit partagé et somme privée", "protège typiquement mises à jour/statistiques, pas une esquisse géométrique de frontière [R12–R14]"),
        ("Reconstruction fédérée de surfaces 3D", "champs/surfaces en fédéré", "vise une scène physique, pas une frontière de décision de classifieur"),
    ], [1900, 2650, 4810], 8.7)

    doc.add_heading("10.2 Résultat de l’audit ciblé", level=2)
    doc.add_paragraph(
        "Les recherches ciblées effectuées sur les expressions “federated signed-distance field”, “federated Eikonal decision boundary reconstruction” et “Hamilton–Jacobi federated classification” n’ont pas identifié, à la date de ce document, une méthode qui réunisse exactement : espace latent commun, descripteurs locaux orientés, reconstruction d’une SDF de décision, régularisation Eikonal, fusion sensible à la masse/topologie, puis protection par DP distribuée compatible SecAgg. Elles ont toutefois identifié plusieurs travaux adjacents ci-dessus. Une absence de résultat de recherche n’est pas une preuve d’absence dans la littérature."
    )
    add_callout(doc, "Formulation prudente pour un article", "À notre connaissance, après recherche dans les bases et mots-clés spécifiés, nous sommes les premiers à formuler la collaboration fédérée comme la reconstruction privée d’un champ de distance signée de décision à partir de contraintes locales orientées. Cette revendication devra être réauditée avant soumission.", "claim")

    doc.add_heading("10.3 Les objections probables d’un reviewer", level=2)
    add_table(doc, ["Objection", "Réponse expérimentale attendue"], [
        ("“C’est de la distillation avec un nouveau vocabulaire.”", "Comparer aux logits sur les mêmes sondes; montrer le gain spécifique des normales/Eikonal/UOT."),
        ("“Vous additionnez quand même des vecteurs.”", "Distinguer primitive SecAgg et objet d’apprentissage; fournir une version MPC comme borne conceptuelle."),
        ("“La métrique latente est arbitraire.”", "Encodeur gelé, tests de voisinage sémantique, sensibilité aux transformations du latent."),
        ("“La DP détruit les détails fins.”", "Courbe ε–Boundary-F1–longueur; esquisses multi-échelles et zones critiques."),
        ("“Les frontières locales sont contradictoires.”", "Confiance, UOT non équilibré, reconstruction robuste, analyse par client et non-IID."),
        ("“SDF n’est pas différentiable partout.”", "Énoncer presque partout; échantillonner hors axe médian et utiliser une régularisation souple."),
        ("“La segmentation est trop spécifique.”", "Présenter le cadre général, mais revendiquer d’abord un résultat solide et falsifiable."),
    ], [3300, 6060], 8.9)

    doc.add_heading("11. Programme théorique", level=1)
    doc.add_heading("11.1 Questions démontrables", level=2)
    add_bullet(doc, "Erreur de projection locale : borner |δ_exact−δ_Taylor| en fonction de la Hessienne de m et de ‖∇m‖.")
    add_bullet(doc, "Stabilité de reconstruction : relier l’erreur de l’esquisse ‖Ũ−U‖ à une erreur de niveau zéro, sous une hypothèse de gradient minimal |∇φ|≥γ près de Γ.")
    add_bullet(doc, "Confidentialité–géométrie : borner l’erreur de moments due au clipping, à la quantification et au bruit distribué.")
    add_bullet(doc, "Communication : établir D, bits/coordonnée et nombre de sondes nécessaires pour une précision de contour donnée.")
    add_bullet(doc, "Robustesse non-IID : caractériser quand une frontière globale majoritaire efface une frontière minoritaire mais critique.")
    doc.add_heading("11.2 Exemple de chaîne de borne", level=2)
    doc.add_paragraph(
        "Si φ⋆ est le champ idéal, φ̂ le champ reconstruit, et si ‖φ̂−φ⋆‖∞≤α dans un tube où ‖∇φ⋆‖≥γ>0, une intuition classique de stabilité des level sets suggère une erreur de localisation de l’ordre de α/γ. Le programme théorique consiste à décomposer α :"
    )
    add_eq(doc, "α ≤ α_projection + α_esquisse + α_quantification + α_DP + α_optimisation.", "24")
    doc.add_paragraph(
        "Cette équation est une feuille de route, pas un théorème déjà prouvé. Chaque terme exige des hypothèses. Elle a néanmoins une grande valeur scientifique : elle sépare les sources d’erreur au lieu d’attribuer toute dégradation à la DP ou à l’hétérogénéité."
    )

    doc.add_heading("12. Ce qu’il faut apprendre — parcours guidé", level=1)
    doc.add_heading("12.1 Socle 1 : calcul différentiel et géométrie implicite", level=2)
    add_bullet(doc, "Gradient, Jacobienne, Hessienne, développement de Taylor, différentiation automatique.")
    add_bullet(doc, "Théorème des fonctions implicites, espaces tangents, normales, courbure.")
    add_bullet(doc, "Distances à un ensemble, projections, reach, axe médian et singularités des SDF.")
    add_bullet(doc, "Exercice maître : dériver (6)–(10), puis vérifier numériquement l’erreur de δ sur une ellipse et un petit réseau.")

    doc.add_heading("12.2 Socle 2 : EDP, level sets et Hamilton–Jacobi", level=2)
    add_bullet(doc, "Équation d’Eikonal, solutions de viscosité, fast marching/fast sweeping.")
    add_bullet(doc, "Méthodes level set, re-distanciation et évolution de fronts [R19].")
    add_bullet(doc, "Lien entre Hamilton–Jacobi, temps d’arrivée et classification sur graphes.")
    add_bullet(doc, "Exercice maître : reconstruire une SDF 2D depuis des points de contour bruités, d’abord sur grille, puis par MLP.")

    doc.add_heading("12.3 Socle 3 : transport optimal et topologie", level=2)
    add_bullet(doc, "Mesures, couplages, Wasserstein, Sinkhorn, biais entropique.")
    add_bullet(doc, "Transport non équilibré, divergences de marginales, géométrie des mesures [R6–R7].")
    add_bullet(doc, "Notions de topologie computationnelle : composantes, trous, homologie persistante — seulement après le prototype UOT.")
    add_bullet(doc, "Exercice maître : construire deux contours proches au sens de Hausdorff mais de longueurs différentes, puis comparer HD, MHD et UOT.")

    doc.add_heading("12.4 Socle 4 : apprentissage fédéré et confidentialité", level=2)
    add_bullet(doc, "FedAvg, non-IID, participation partielle, distillation, personnalisation et agrégation robuste [R1–R2].")
    add_bullet(doc, "Définition DP, adjacence, sensibilité, clipping, mécanisme gaussien, RDP et composition [R9–R10].")
    add_bullet(doc, "Agrégation sécurisée, quantification modulaire, dropouts, collusion, gaussienne discrète, Skellam, PBM [R11–R14].")
    add_bullet(doc, "Exercice maître : implémenter une somme de sketches avec clipping et vérifier empiriquement la variance après différents taux d’abandon.")

    doc.add_heading("12.5 Socle 5 : expérimentation scientifique", level=2)
    add_bullet(doc, "Segmentation, pertes de région/frontière, calibration, statistiques par seeds et intervalles de confiance.")
    add_bullet(doc, "Threat modeling, attaques de copie/extraction, tests d’inférence et limites honnêtes.")
    add_bullet(doc, "Mesure énergie/communication du dépôt : joules, octets, coût jusqu’à une qualité cible, pas seulement accuracy finale.")
    add_table(doc, ["Période", "Objectif", "Livrable"], [
        ("Semaines 1–3", "géométrie + SDF 2D", "notebook analytique et reconstruction d’une ellipse"),
        ("Semaines 4–6", "SDF neuronale + segmentation", "baseline centralisée avec métriques de contour"),
        ("Semaines 7–9", "FedBCR non privé", "pipeline multi-client et ablations normales/Eikonal"),
        ("Semaines 10–12", "esquisse + SecAgg simulé", "équivalence point/sketch et coût communication"),
        ("Semaines 13–15", "DP distribuée", "comptable, dropouts, courbes ε–utilité"),
        ("Semaine 16", "audit et décision", "rapport go/no-go, figures reproductibles, claim révisé"),
    ], [1500, 3220, 4640], 9.2)

    doc.add_heading("13. Feuille de route expérimentale", level=1)
    doc.add_heading("13.1 Prototype A — prouver la géométrie", level=2)
    add_number(doc, "Créer des clients synthétiques avec frontières connues et hétérogènes.")
    add_number(doc, "Comparer moyenne de paramètres, distillation de logits et FedBCR sur la reconstruction du niveau zéro.")
    add_number(doc, "Tester l’erreur de Taylor, Newton, les normales et la sensibilité au bruit.")
    add_number(doc, "Critère go : FedBCR doit reconstruire une géométrie correcte lorsque les paramétrisations locales sont incompatibles.")
    doc.add_heading("13.2 Prototype B — segmentation industrielle non privée", level=2)
    add_number(doc, "Encodeur gelé et tête légère; partition par domaine industriel.")
    add_number(doc, "Sondes publiques/synthétiques; esquisse sans bruit.")
    add_number(doc, "Comparer MHD, UOT, HD95 et longueur; sélectionner la fonction de fidélité.")
    add_number(doc, "Critère go : gain de frontière statistiquement robuste sans perte prohibitive de Dice/IoU.")
    doc.add_heading("13.3 Prototype C — FedBCR-DP", level=2)
    add_number(doc, "Quantification discrète et SecAgg simulé avec dropouts.")
    add_number(doc, "DP distribuée avec unité protégée explicite; comptabilité multi-round.")
    add_number(doc, "Mesurer attaques de copie et extraction avant/après protection.")
    add_number(doc, "Critère go : avantage d’attaque réduit et frontière utile dans un budget ε crédible.")

    doc.add_heading("14. Risques, limites et critères d’abandon", level=1)
    add_table(doc, ["Risque", "Signal d’alerte", "Réponse / kill criterion"], [
        ("Repère latent non aligné", "normales incohérentes entre clients", "geler E; abandonner la version multi-encodeur si l’alignement échoue"),
        ("Malédiction de dimension", "D et sondes explosent", "réduire d, base multi-échelle creuse; abandonner grille dense"),
        ("Frontières minoritaires effacées", "rappel défaut rare chute", "contraintes de couverture/robustesse; priorité aux zones critiques"),
        ("Bruit DP destructeur", "Boundary-F1 s’effondre pour ε acceptable", "réduire dimension/fréquence; changer l’unité protégée avec justification"),
        ("Coût supérieur à FedAvg", "énergie/latence prohibitives", "distillation compacte, moins de sondes; rapporter honnêtement le domaine utile"),
        ("Aucune différence aux logits", "ablation normale/Eikonal n’apporte rien", "reformuler la contribution; ne pas conserver un vocabulaire géométrique inutile"),
        ("Antériorité trouvée", "méthode quasi identique publiée", "positionner comme extension privée/UOT/énergie ou changer de question"),
    ], [2200, 2650, 4510], 8.8)
    add_callout(doc, "Le meilleur résultat négatif possible", "Si la distillation de logits sur sondes publiques égale FedBCR à coût inférieur, le projet aura quand même répondu à une question scientifique importante : les dérivées géométriques ne sont pas nécessaires dans ce régime. Le protocole expérimental doit permettre cette conclusion.", "warning")

    doc.add_heading("15. Contribution proposée, en couches", level=1)
    add_table(doc, ["Niveau", "Contribution", "Preuve attendue"], [
        ("C1 — Paradigme", "fédérer une frontière implicite plutôt qu’un modèle", "définition formelle + cas de paramétrisations hétérogènes"),
        ("C2 — Message", "descripteurs orientés marge–normale–projection avec validité", "dérivation, erreur locale et ablations"),
        ("C3 — Fusion", "SDF/Eikonal avec fidélité UOT et contraintes industrielles", "gain sur défauts fins, longueur et topologie"),
        ("C4 — Confidentialité", "sketch additif, SecAgg et bruit distribué", "preuve DP, dropouts/collusion et attaques"),
        ("C5 — Efficacité", "compression géométrique et métriques énergie", "octets, joules et qualité par coût"),
    ], [1500, 4430, 3430], 9.0)
    doc.add_paragraph(
        "Le cœur intellectuel est C1+C2. C3 donne une forme mathématique adaptée à l’industrie. C4 rend l’idée crédible du point de vue sécurité. C5 l’inscrit naturellement dans un programme de thèse sur l’apprentissage fédéré économe en énergie. Toutes les couches ne doivent pas être revendiquées comme nouveautés indépendantes : elles forment une contribution systémique."
    )

    doc.add_heading("16. Conclusion pédagogique", level=1)
    doc.add_paragraph(
        "FedBCR invite à regarder l’apprentissage fédéré non comme une mécanique de moyenne, mais comme un problème inverse : plusieurs modèles privés observent chacun une partie du monde; le serveur doit reconstruire la géométrie d’une décision collective. La marge dit de quel côté on se trouve. Son gradient donne la normale. Taylor donne un premier pas vers la frontière. La SDF transforme une surface en fonction continue. Eikonal discipline cette fonction pour qu’elle ressemble à une distance. UOT prend au sérieux la masse et la longueur des contours. La DP distribuée limite l’influence observable de chaque client, tandis que SecAgg masque les messages individuels."
    )
    doc.add_paragraph(
        "La force de l’idée ne viendra pas de la sophistication des mots, mais de trois démonstrations : elle reconstruit mieux une frontière utile que les alternatives; elle protège effectivement les participants; et son coût de communication/énergie reste justifié. Le premier terrain — segmentation binaire de défauts industriels — est assez concret pour faire échouer l’idée si elle est fausse et assez riche pour révéler son intérêt si elle est juste. C’est exactement le type d’expérience qu’une contribution doctorale doit rechercher."
    )

    doc.add_page_break()
    doc.add_heading("Annexe A — Dérivation complète de la projection", level=1)
    doc.add_paragraph(
        "On cherche le point le plus proche de z_q sur m(z)=0. Le problème exact est min_b ½‖b−z_q‖² sous m(b)=0. Le Lagrangien est ℒ(b,λ)=½‖b−z_q‖²+λm(b). Les conditions du premier ordre sont b−z_q+λ∇m(b)=0 et m(b)=0. Elles montrent que le déplacement optimal est normal à la surface au point b. Comme b est inconnu, on linéarise m autour de z_q :"
    )
    add_eq(doc, "m(b) ≈ m(z_q)+∇m(z_q)ᵀ(b−z_q)=0.", "A1")
    doc.add_paragraph("Posons d=b−z_q et g=∇m(z_q). Le problème linéarisé est min_d ½‖d‖² sous gᵀd=−m. Son Lagrangien donne d+λg=0, donc d=−λg. En réinjectant :")
    add_eq(doc, "−λ‖g‖²=−m  ⇒  λ=m/‖g‖²  ⇒  d=−m g/‖g‖².", "A2")
    doc.add_paragraph("En écrivant n=g/‖g‖ et d=δn, on retrouve :")
    add_eq(doc, "δ=−m/‖g‖,  et  b≈z_q−m(z_q)∇m(z_q)/‖∇m(z_q)‖².", "A3")
    doc.add_paragraph(
        "La formule (9) stabilise δ directement; la formule (10) stabilise le dénominateur quadratique. Elles ne sont pas exactement identiques lorsque η>0. La seconde correspond plus naturellement à la projection de Newton et doit être privilégiée pour l’implémentation."
    )

    doc.add_heading("Annexe B — Exemple de calibration DP", level=1)
    doc.add_paragraph(
        "Exemple illustratif uniquement : K=50 clients sélectionnés, h_min=30 contributions honnêtes garanties après collusion et abandon, clipping C=1, adjacence ajout/retrait donc Δ≤1. Si le mécanisme exige une variance agrégée v_target par coordonnée, on choisit des parts indépendantes dont la convolution atteint au moins v_target avec seulement h_min participants. Pour une approximation gaussienne v_target=σ²Δ² et τ²=v_target/h_min. Avec davantage de clients honnêtes, la variance devient plus élevée que le minimum, donc la confidentialité reste conservatrice mais l’utilité baisse légèrement."
    )
    add_callout(doc, "Ne pas publier un ε improvisé", "Le vrai ε dépend du mécanisme discret, de la quantification, du module, du sous-échantillonnage, du nombre de rounds, de δ et du modèle de collusion. Utiliser l’analyse officielle du mécanisme et un comptable reproductible.", "privacy")

    doc.add_page_break()
    doc.add_heading("Annexe C — Checklist avant expérience", level=1)
    checks = [
        "Unité de confidentialité et relation de voisinage écrites noir sur blanc.",
        "Encodeur commun, métrique latente et protocole de sondes figés avant test.",
        "Seuils de validité des normales et projections pré-enregistrés.",
        "Aucune sonde ni ancre issue d’une image privée brute.",
        "Dimension, clipping, quantification, module, h_min et modèle de collusion documentés.",
        "Baselines utilisant exactement les mêmes données publiques et budget de communication.",
        "Métriques de région, frontière, masse, topologie, énergie et confidentialité rapportées.",
        "Au moins cinq seeds ou justification statistique adaptée; intervalles de confiance.",
        "Ablations normales, Eikonal, UOT, sketch et DP.",
        "Audit d’antériorité relancé juste avant soumission.",
    ]
    for item in checks:
        add_bullet(doc, "☐ " + item)

    doc.add_heading("Références", level=1)
    references = [
        ("McMahan, B. et al. (2017). Communication-Efficient Learning of Deep Networks from Decentralized Data. AISTATS, PMLR 54.", "https://proceedings.mlr.press/v54/mcmahan17a.html"),
        ("Kairouz, P. et al. (2021). Advances and Open Problems in Federated Learning. Foundations and Trends in Machine Learning, 14(1–2).", "https://arxiv.org/abs/1912.04977"),
        ("Park, J. J. et al. (2019). DeepSDF: Learning Continuous Signed Distance Functions for Shape Representation. CVPR.", "https://openaccess.thecvf.com/content_CVPR_2019/html/Park_DeepSDF_Learning_Continuous_Signed_Distance_Functions_for_Shape_Representation_CVPR_2019_paper.html"),
        ("Gropp, A. et al. (2020). Implicit Geometric Regularization for Learning Shapes. ICML, PMLR 119.", "https://proceedings.mlr.press/v119/gropp20a.html"),
        ("Dubuisson, M.-P., Jain, A. K. (1994). A Modified Hausdorff Distance for Object Matching. ICPR.", "https://doi.org/10.1109/ICPR.1994.576361"),
        ("Peyré, G., Cuturi, M. (2019). Computational Optimal Transport. Foundations and Trends in Machine Learning.", "https://arxiv.org/abs/1803.00567"),
        ("Chizat, L. et al. (2018). Unbalanced Optimal Transport: Dynamic and Kantorovich Formulations. Journal of Functional Analysis.", "https://arxiv.org/abs/1508.05216"),
        ("Yousefzadeh, R. (2021). Decision Boundaries and Convex Hulls in the Feature Space that Deep Learning Functions Learn from Images. arXiv:2103.00695.", "https://arxiv.org/abs/2103.00695"),
        ("Dwork, C., Roth, A. (2014). The Algorithmic Foundations of Differential Privacy. Foundations and Trends in Theoretical Computer Science.", "https://www.cis.upenn.edu/~aaroth/Papers/privacybook.pdf"),
        ("Abadi, M. et al. (2016). Deep Learning with Differential Privacy. CCS.", "https://arxiv.org/abs/1607.00133"),
        ("Bonawitz, K. et al. (2017). Practical Secure Aggregation for Privacy-Preserving Machine Learning. CCS.", "https://eprint.iacr.org/2017/281.pdf"),
        ("Kairouz, P., Liu, Z., Steinke, T. (2021). The Distributed Discrete Gaussian Mechanism for Federated Learning with Secure Aggregation. ICML, PMLR 139.", "https://proceedings.mlr.press/v139/kairouz21a.html"),
        ("Agarwal, N. et al. (2021). The Skellam Mechanism for Differentially Private Federated Learning. NeurIPS 34.", "https://proceedings.neurips.cc/paper/2021/hash/285baacbdf8fda1de94b19282acd23e2-Abstract.html"),
        ("Chen, W.-N. et al. (2022). Poisson Binomial Mechanism for Unbiased Federated Learning with Secure Aggregation. ICML, PMLR 162.", "https://proceedings.mlr.press/v162/chen22s.html"),
        ("Calder, J., Ettehad, M. (2022). Hamilton–Jacobi Equations on Graphs with Applications to Semi-Supervised Learning and Data Depth. JMLR 23.", "https://www.jmlr.org/papers/v23/22-0293.html"),
        ("Fed-GDBD (2025). Federated Global Decision Boundary Distillation. Discover Computing.", "https://link.springer.com/article/10.1007/s44443-025-00097-0"),
        ("Kervadec, H. et al. (2019). Boundary Loss for Highly Unbalanced Segmentation. MIDL / Medical Image Analysis.", "https://arxiv.org/abs/1812.07032"),
        ("Karimi, D., Salcudean, S. E. (2020). Reducing the Hausdorff Distance in Medical Image Segmentation with Convolutional Neural Networks. IEEE TMI.", "https://pubmed.ncbi.nlm.nih.gov/31329113/"),
        ("Osher, S., Sethian, J. A. (1988). Fronts Propagating with Curvature-Dependent Speed: Algorithms Based on Hamilton–Jacobi Formulations. Journal of Computational Physics.", "https://doi.org/10.1016/0021-9991(88)90002-2"),
        ("Chen, W.-N. et al. (2022). The Fundamental Price of Secure Aggregation in Differentially Private Federated Learning. ICML, PMLR 162.", "https://proceedings.mlr.press/v162/chen22c.html"),
        ("Chen, W.-N. et al. (2023). The Communication Cost of Security and Privacy in Federated Frequency Estimation. AISTATS, PMLR 206.", "https://proceedings.mlr.press/v206/chen23e.html"),
        ("Fainstein, M. et al. (2024). DUDF: Differentiable Unsigned Distance Fields with Hyperbolic Scaling. CVPR.", "https://openaccess.thecvf.com/content/CVPR2024/html/Fainstein_DUDF_Differentiable_Unsigned_Distance_Fields_with_Hyperbolic_Scaling_CVPR_2024_paper.html"),
        ("Liu, L. et al. (2024). SurroundSDF: Implicit 3D Scene Understanding Based on Signed Distance Field. CVPR.", "https://openaccess.thecvf.com/content/CVPR2024/html/Liu_SurroundSDF_Implicit_3D_Scene_Understanding_Based_on_Signed_Distance_Field_CVPR_2024_paper.html"),
    ]
    for idx, (citation, url) in enumerate(references, 1):
        add_ref(doc, idx, citation, url)

    doc.add_paragraph()
    p = doc.add_paragraph(style="Small Note")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Fin du document — FedBCR-DP, version conceptuelle du 17 juillet 2026.")

    props = doc.core_properties
    props.title = "FedBCR-DP — Apprendre une frontière globale sans moyenner les modèles"
    props.subject = "Cadre de recherche en apprentissage fédéré, SDF/Eikonal, UOT et confidentialité différentielle distribuée"
    props.author = "Document de travail doctoral"
    props.keywords = "federated learning; decision boundary; signed distance field; Eikonal; distributed differential privacy; secure aggregation"
    props.comments = "Proposition conceptuelle à valider expérimentalement et bibliographiquement."
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_document())
